import os
from django.test import TestCase
from django.core.files.uploadedfile import SimpleUploadedFile
from django.urls import reverse
from rest_framework.test import APIClient
from docx import Document
from templates.models import Template
from resumes.models import Resume, ResumeVersion, ResumeSection
from .models import GeneratedDocument
from .services import generate_docx_for_resume

from django.contrib.auth import get_user_model
User = get_user_model()

class DocumentGenerationTests(TestCase):
    def setUp(self):
        self.client = APIClient()
        self.user = User.objects.create_user(username='testuser', password='password123')
        self.client.force_authenticate(user=self.user)
        
        # Create a dummy docx template for testing
        self.template_path = 'test_template.docx'
        doc = Document()
        doc.add_paragraph('Name: {{NAME}}')
        doc.add_paragraph('Experience: {{EXPERIENCE}}')
        doc.save(self.template_path)
        
        with open(self.template_path, 'rb') as f:
            self.template = Template.objects.create(
                name="Test Template",
                file=SimpleUploadedFile("test_template.docx", f.read(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            )
            
        # Create Resume and Sections
        self.resume = Resume.objects.create(title="My Resume", user=self.user)
        self.version = ResumeVersion.objects.create(resume=self.resume, version_name="v1")
        
        ResumeSection.objects.create(
            version=self.version,
            section_type='personal',
            content={'first_name': 'John', 'last_name': 'Doe'},
            order=0
        )
        ResumeSection.objects.create(
            version=self.version,
            section_type='experience',
            content={
                'job_title': 'Software Engineer',
                'company_name': 'Tech Corp',
                'start_date': '2020',
                'end_date': '2022',
                'description': 'Did some coding.'
            },
            order=1
        )
        ResumeSection.objects.create(
            version=self.version,
            section_type='experience',
            content={
                'job_title': 'Senior Engineer',
                'company_name': 'Big Tech',
                'start_date': '2022',
                'end_date': 'Present',
                'description': 'Did more coding.'
            },
            order=2
        )

    def tearDown(self):
        try:
            if os.path.exists(self.template_path):
                os.remove(self.template_path)
        except OSError:
            pass
            
        # Clean up generated files
        for gd in GeneratedDocument.objects.all():
            try:
                if gd.file and os.path.exists(gd.file.path):
                    os.remove(gd.file.path)
            except OSError:
                pass
                
        try:
            if self.template.file and os.path.exists(self.template.file.path):
                os.remove(self.template.file.path)
        except OSError:
            pass

    def test_generate_docx_service(self):
        gen_doc = generate_docx_for_resume(self.version.id, self.template.id)
        
        self.assertIsNotNone(gen_doc)
        self.assertEqual(gen_doc.resume_version, self.version)
        self.assertEqual(gen_doc.template, self.template)
        
        # Verify generated file exists and can be opened
        self.assertTrue(os.path.exists(gen_doc.file.path))
        doc = Document(gen_doc.file.path)
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        self.assertIn('Name: John Doe', full_text)
        self.assertIn('Software Engineer at Tech Corp (2020 - 2022)\nDid some coding.', full_text)
        self.assertIn('Senior Engineer at Big Tech (2022 - Present)\nDid more coding.', full_text)

    def test_generate_document_api(self):
        url = reverse('generate-document')
        response = self.client.post(url, {
            'resume_version_id': self.version.id,
            'template_id': self.template.id
        })
        self.assertEqual(response.status_code, 201)
        self.assertIn('download_url', response.data)
        
        # Test download API
        download_url = response.data['download_url']
        dl_response = self.client.get(download_url)
        self.assertEqual(dl_response.status_code, 200)
        self.assertIn(f'filename="resume_{self.version.id}_{self.template.id}', dl_response.get('Content-Disposition'))

    def test_upload_pdf_document_api(self):
        url = reverse('upload-pdf')
        
        # Create a dummy PDF file
        pdf_content = b'%PDF-1.4\n%EOF'
        pdf_file = SimpleUploadedFile("dummy.pdf", pdf_content, content_type="application/pdf")
        
        response = self.client.post(url, {
            'resume_version_id': self.version.id,
            'template_id': self.template.id,
            'file': pdf_file
        }, format='multipart')
        
        self.assertEqual(response.status_code, 201)
        self.assertIn('download_url', response.data)
        self.assertEqual(response.data['format'], 'pdf')
        
        # Test download API for PDF
        download_url = response.data['download_url']
        dl_response = self.client.get(download_url)
        self.assertEqual(dl_response.status_code, 200)
        self.assertIn(f'filename="resume_{self.version.id}_{self.template.id}', dl_response.get('Content-Disposition'))

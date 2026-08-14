import io
import re
from docx import Document
from django.core.files.base import ContentFile
from .models import GeneratedDocument
from templates.models import Template
from resumes.models import ResumeVersion

def _resolve_placeholder(placeholder, resume_version):
    """
    Given a placeholder like 'NAME', 'EXPERIENCE', resolves it from resume_version.
    Returns a string representation of the data.
    """
    if not resume_version:
        return ""
    
    # sections = resume_version.sections.all()
    sections = resume_version.sections.all()
    
    def get_section_data(section_type):
        sec = next((s for s in sections if s.section_type == section_type), None)
        return sec.content if sec else {}
    
    def get_list_section_data(section_type):
        return [s.content for s in sections if s.section_type == section_type]
    
    if placeholder == 'NAME':
        data = get_section_data('personal')
        return f"{data.get('first_name', '')} {data.get('last_name', '')}".strip()
    
    if placeholder == 'EMAIL':
        return get_section_data('personal').get('email', '')
        
    if placeholder == 'PHONE':
        return get_section_data('personal').get('phone', '')
        
    if placeholder == 'LOCATION':
        return get_section_data('personal').get('location', '')
        
    if placeholder == 'SUMMARY':
        return get_section_data('summary').get('text', '')
        
    if placeholder == 'EXPERIENCE':
        exps = get_list_section_data('experience')
        lines = []
        for exp in exps:
            title = exp.get('job_title', '')
            company = exp.get('company_name', '')
            start = exp.get('start_date', '')
            end = exp.get('end_date', 'Present')
            desc = exp.get('description', '')
            
            header = []
            if title: header.append(title)
            if company: header.append(f"at {company}")
            if start or end: header.append(f"({start} - {end})")
            
            if header: lines.append(" ".join(header))
            if desc: lines.append(desc)
            lines.append("")
        return "\n".join(lines).strip()
        
    if placeholder == 'EDUCATION':
        edus = get_list_section_data('education')
        lines = []
        for ed in edus:
            degree = ed.get('degree', '')
            school = ed.get('school_name', '')
            start = ed.get('start_date', '')
            end = ed.get('end_date', 'Present')
            
            header = []
            if degree: header.append(degree)
            if school: header.append(f"from {school}")
            if start or end: header.append(f"({start} - {end})")
            
            if header: lines.append(" ".join(header))
        return "\n".join(lines).strip()
        
    if placeholder == 'PROJECTS':
        projs = get_list_section_data('projects')
        lines = []
        for pr in projs:
            name = pr.get('name', '')
            desc = pr.get('description', '')
            if name: lines.append(name)
            if desc: lines.append(desc)
            lines.append("")
        return "\n".join(lines).strip()
        
    if placeholder == 'SKILLS':
        skills_secs = get_list_section_data('skills')
        lines = []
        for s in skills_secs:
            if 'name' in s:
                lines.append(s['name'])
            elif 'items' in s:
                if isinstance(s['items'], list):
                    lines.extend([str(x) for x in s['items']])
                else:
                    lines.append(str(s['items']))
            elif 'text' in s:
                lines.append(s['text'])
        return ", ".join(lines)
        
    if placeholder == 'CERTIFICATIONS':
        certs = get_list_section_data('certifications')
        lines = []
        for c in certs:
            name = c.get('name', '')
            if name: lines.append(name)
        return ", ".join(lines)

    return ""

def _resolve_placeholder_from_data(placeholder, resume_data):
    if not resume_data:
        return ""
        
    if placeholder == 'NAME':
        return f"{resume_data.get('personalInfo', {}).get('firstName', '')} {resume_data.get('personalInfo', {}).get('lastName', '')}".strip()
    
    if placeholder == 'EMAIL':
        return resume_data.get('personalInfo', {}).get('email', '')
        
    if placeholder == 'PHONE':
        return resume_data.get('personalInfo', {}).get('phone', '')
        
    if placeholder == 'LOCATION':
        return resume_data.get('personalInfo', {}).get('location', '')
        
    if placeholder == 'SUMMARY':
        return resume_data.get('personalInfo', {}).get('summary', '')
        
    if placeholder == 'EXPERIENCE':
        exps = resume_data.get('experience', [])
        lines = []
        for exp in exps:
            title = exp.get('title', '')
            company = exp.get('company', '')
            start = exp.get('startDate', '')
            end = exp.get('endDate', 'Present')
            desc = exp.get('description', '')
            
            header = []
            if title: header.append(title)
            if company: header.append(f"at {company}")
            if start or end: header.append(f"({start} - {end})")
            
            if header: lines.append(" ".join(header))
            if desc: lines.append(desc)
            lines.append("")
        return "\n".join(lines).strip()
        
    if placeholder == 'EDUCATION':
        edus = resume_data.get('education', [])
        lines = []
        for ed in edus:
            degree = ed.get('degree', '')
            school = ed.get('school', '')
            start = ed.get('startDate', '')
            end = ed.get('endDate', 'Present')
            
            header = []
            if degree: header.append(degree)
            if school: header.append(f"from {school}")
            if start or end: header.append(f"({start} - {end})")
            
            if header: lines.append(" ".join(header))
        return "\n".join(lines).strip()
        
    if placeholder == 'PROJECTS':
        projs = resume_data.get('projects', [])
        lines = []
        for pr in projs:
            name = pr.get('name', '')
            desc = pr.get('description', '')
            if name: lines.append(name)
            if desc: lines.append(desc)
            lines.append("")
        return "\n".join(lines).strip()
        
    if placeholder == 'SKILLS':
        skills = resume_data.get('skills', [])
        lines = []
        for s in skills:
            if isinstance(s, dict) and 'name' in s:
                lines.append(s['name'])
            elif isinstance(s, str):
                lines.append(s)
        return ", ".join(lines)
        
    if placeholder == 'CERTIFICATIONS':
        certs = resume_data.get('certifications', [])
        lines = []
        for c in certs:
            if isinstance(c, dict) and 'name' in c:
                lines.append(c['name'])
            elif isinstance(c, str):
                lines.append(c)
        return ", ".join(lines)

    return ""

def _replace_text_in_paragraphs_from_data(paragraphs, resume_data):
    pattern = re.compile(r'\{\{([A-Z_]+)\}\}')
    for para in paragraphs:
        if '{{' in para.text:
            for run in para.runs:
                matches = pattern.findall(run.text)
                for ph in matches:
                    val = _resolve_placeholder_from_data(ph, resume_data)
                    run.text = run.text.replace(f"{{{{{ph}}}}}", val)
            
            if '{{' in para.text and '}}' in para.text:
                full_text = para.text
                matches = pattern.findall(full_text)
                if matches:
                    for ph in matches:
                        val = _resolve_placeholder_from_data(ph, resume_data)
                        full_text = full_text.replace(f"{{{{{ph}}}}}", val)
                    
                    if para.runs:
                        para.runs[0].text = full_text
                        for i in range(1, len(para.runs)):
                            para.runs[i].text = ""

def generate_docx_for_resume(resume_version_id, template_id, resume_data=None, user=None):
    """
    Generates a DOCX file for the given resume_version and template.
    If resume_data is provided, it skips DB loading and uses the data dictionary directly.
    Saves it to a GeneratedDocument and returns it.
    """
    try:
        if isinstance(template_id, int) or str(template_id).isdigit():
            template = Template.objects.get(id=int(template_id))
        else:
            template = Template.objects.filter(name__iexact=str(template_id)).first()
            if not template:
                template = Template.objects.first()
    except Template.DoesNotExist:
        template = Template.objects.first()

    doc = Document(template.file)
    
    if resume_data:
        # Client-side data flow
        _replace_text_in_paragraphs_from_data(doc.paragraphs, resume_data)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _replace_text_in_paragraphs_from_data(cell.paragraphs, resume_data)
                    
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        gen_doc = GeneratedDocument(
            template=template,
            format='docx',
            user=user
        )
        # Note: resume_version might be null if not saved
        if resume_version_id:
            try:
                gen_doc.resume_version = ResumeVersion.objects.get(id=resume_version_id)
            except ResumeVersion.DoesNotExist:
                pass
                
        file_name = f"resume_export_{template.id}.docx"
        gen_doc.file.save(file_name, ContentFile(buffer.read()))
        gen_doc.save()
        return gen_doc
    else:
        # Original flow using DB
        resume_version = ResumeVersion.objects.get(id=resume_version_id)
        _replace_text_in_paragraphs(doc.paragraphs, resume_version)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _replace_text_in_paragraphs(cell.paragraphs, resume_version)
                    
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        file_name = f"resume_{resume_version.id}_{template.id}.docx"
        gen_doc = GeneratedDocument(
            resume_version=resume_version,
            template=template,
            format='docx',
            user=user
        )
        gen_doc.file.save(file_name, ContentFile(buffer.read()))
        gen_doc.save()
        return gen_doc

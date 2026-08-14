from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from rest_framework.permissions import IsAuthenticated
from rest_framework.parsers import MultiPartParser
from .services.generation import generate_resume_content
from .services.extraction import extract_resume_data
from .services.file_parsing import extract_text_from_pdf, extract_text_from_docx
from .services.job_tailoring import tailor_resume_for_job
from .services.interview_prep import generate_interview_questions
from .services.matching import calculate_job_match
from .services.ats_analysis import analyze_ats_compatibility
from .providers.openai_provider import AIProviderValidationException
import logging
import json
from resumes.models import ResumeVersion
from job_matching.models import JobDescription

logger = logging.getLogger(__name__)

class GenerateResumeView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        resume_data = request.data.get('resume_data')
        template_info = request.data.get('template_info')

        if not resume_data:
            return Response({"error": "resume_data is required"}, status=status.HTTP_400_BAD_REQUEST)

        try:
            generated_content = generate_resume_content(resume_data, template_info)
            return Response(generated_content.model_dump(), status=status.HTTP_200_OK)
        except AIProviderValidationException as e:
            logger.error(f"AI Validation Error: {e}")
            return Response({"error": "AI response failed validation", "details": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        except Exception as e:
            logger.error(f"AI Generation Error: {e}")
            return Response({"error": "An error occurred during AI generation", "details": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class ExtractResumeView(APIView):
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser]

    def post(self, request, *args, **kwargs):
        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({"error": "No file uploaded"}, status=status.HTTP_400_BAD_REQUEST)

        # Validate size (max 5MB)
        if file_obj.size > 5 * 1024 * 1024:
            return Response({"error": "File size exceeds 5MB limit"}, status=status.HTTP_400_BAD_REQUEST)

        # Validate and parse file type
        filename = file_obj.name.lower()
        file_bytes = file_obj.read()
        
        if not file_bytes:
            return Response({"error": "Empty file uploaded"}, status=status.HTTP_400_BAD_REQUEST)

        try:
            if filename.endswith('.pdf'):
                raw_text = extract_text_from_pdf(file_bytes)
            elif filename.endswith('.docx'):
                raw_text = extract_text_from_docx(file_bytes)
            else:
                return Response({"error": "Unsupported file format. Please upload PDF or DOCX."}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            logger.error(f"File parsing error: {e}")
            return Response({"error": "Failed to read the file. It may be corrupted."}, status=status.HTTP_400_BAD_REQUEST)
            
        if not raw_text.strip():
            return Response({"error": "No text could be extracted from the document."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            extracted_data = extract_resume_data(raw_text)
            return Response(extracted_data.model_dump(), status=status.HTTP_200_OK)
        except AIProviderValidationException as e:
            logger.error(f"AI Validation Error during extraction: {e}")
            return Response({"error": "AI failed to extract valid structured data.", "details": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        except Exception as e:
            logger.error(f"AI Extraction Error: {e}")
            return Response({"error": "An error occurred during AI extraction.", "details": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class OptimizeResumeView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        resume_version_id = request.data.get('resume_version_id')
        custom_resume_text = request.data.get('custom_resume_text')
        job_description_id = request.data.get('job_description_id')
        custom_job_text = request.data.get('job_description_text')

        if not (resume_version_id or custom_resume_text):
            return Response({"error": "resume_version_id or custom_resume_text is required"}, status=status.HTTP_400_BAD_REQUEST)
            
        if not (job_description_id or custom_job_text):
            return Response({"error": "job_description_id or job_description_text is required"}, status=status.HTTP_400_BAD_REQUEST)

        resume_text = custom_resume_text
        if resume_version_id:
            try:
                resume_version = ResumeVersion.objects.get(id=resume_version_id, resume__user=request.user)
                resume_text = ""
                for section in resume_version.sections.all():
                    resume_text += f"{section.title}:\n{json.dumps(section.content)}\n\n"
            except ResumeVersion.DoesNotExist:
                return Response({"error": "Resume not found"}, status=status.HTTP_404_NOT_FOUND)

        job_text = custom_job_text
        if job_description_id:
            try:
                job = JobDescription.objects.get(id=job_description_id, owner=request.user)
                job_text = job.description_text
            except JobDescription.DoesNotExist:
                return Response({"error": "Job not found"}, status=status.HTTP_404_NOT_FOUND)

        try:
            optimized_data = tailor_resume_for_job(resume_text, job_text)
            return Response(optimized_data.model_dump(), status=status.HTTP_200_OK)
        except Exception as e:
            logger.error(f"AI Optimization Error: {e}")
            return Response({"error": "Failed to optimize resume", "details": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class GenerateInterviewPrepView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        resume_version_id = request.data.get('resume_version_id')
        custom_resume_text = request.data.get('custom_resume_text')
        job_description_id = request.data.get('job_description_id')
        custom_job_text = request.data.get('job_description_text')

        if not (resume_version_id or custom_resume_text):
            return Response({"error": "resume_version_id or custom_resume_text is required"}, status=status.HTTP_400_BAD_REQUEST)

        resume_text = custom_resume_text
        if resume_version_id:
            try:
                resume_version = ResumeVersion.objects.get(id=resume_version_id, resume__user=request.user)
                resume_text = ""
                for section in resume_version.sections.all():
                    resume_text += f"{section.title}:\n{json.dumps(section.content)}\n\n"
            except ResumeVersion.DoesNotExist:
                return Response({"error": "Resume version not found"}, status=status.HTTP_404_NOT_FOUND)
            
        job_text = custom_job_text
        if job_description_id:
            try:
                job = JobDescription.objects.get(id=job_description_id, owner=request.user)
                job_text = job.description_text
            except JobDescription.DoesNotExist:
                return Response({"error": "Job description not found"}, status=status.HTTP_404_NOT_FOUND)

        try:
            prep_data = generate_interview_questions(resume_text, job_text)
            return Response(prep_data.model_dump(), status=status.HTTP_200_OK)
        except Exception as e:
            logger.error(f"AI Interview Prep Error: {e}")
            return Response({"error": "Failed to generate interview questions", "details": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class ParseTextView(APIView):
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser]

    def post(self, request, *args, **kwargs):
        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({"error": "No file uploaded"}, status=status.HTTP_400_BAD_REQUEST)

        if file_obj.size > 5 * 1024 * 1024:
            return Response({"error": "File size exceeds 5MB limit"}, status=status.HTTP_400_BAD_REQUEST)

        filename = file_obj.name.lower()
        file_bytes = file_obj.read()
        
        if not file_bytes:
            return Response({"error": "Empty file uploaded"}, status=status.HTTP_400_BAD_REQUEST)

        try:
            if filename.endswith('.pdf'):
                raw_text = extract_text_from_pdf(file_bytes)
            elif filename.endswith('.docx'):
                raw_text = extract_text_from_docx(file_bytes)
            else:
                return Response({"error": "Unsupported file format. Please upload PDF or DOCX."}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            logger.error(f"File parsing error: {e}")
            return Response({"error": "Failed to read the file. It may be corrupted."}, status=status.HTTP_400_BAD_REQUEST)
            
        if not raw_text.strip():
            return Response({"error": "No text could be extracted from the document."}, status=status.HTTP_400_BAD_REQUEST)
            
        return Response({"text": raw_text}, status=status.HTTP_200_OK)

class MatchResumeView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        custom_resume_text = request.data.get('custom_resume_text')
        job_description_text = request.data.get('job_description_text')

        if not custom_resume_text or not job_description_text:
            return Response({"error": "custom_resume_text and job_description_text are required"}, status=status.HTTP_400_BAD_REQUEST)

        try:
            match_data = calculate_job_match(custom_resume_text, job_description_text)
            return Response(match_data.model_dump(), status=status.HTTP_200_OK)
        except Exception as e:
            logger.error(f"AI Matching Error: {e}")
            return Response({"error": "Failed to calculate job match", "details": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class AtsAnalyzerView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        custom_resume_text = request.data.get('custom_resume_text')
        job_description_text = request.data.get('job_description_text', "")

        if not custom_resume_text:
            return Response({"error": "custom_resume_text is required"}, status=status.HTTP_400_BAD_REQUEST)
            
        try:
            try:
                import json
                structured_data_dict = json.loads(custom_resume_text)
            except Exception:
                from .services.extraction import extract_resume_data
                structured_data = extract_resume_data(custom_resume_text)
                structured_data_dict = structured_data.model_dump()
            
            ats_data = analyze_ats_compatibility(structured_data_dict, job_description_text)
            # ats_data is a dict returned by analyze_ats_compatibility
            return Response(ats_data, status=status.HTTP_200_OK)
        except Exception as e:
            logger.error(f"AI ATS Analysis Error: {e}")
            return Response({"error": "Failed to analyze ATS compatibility", "details": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

from pydantic import BaseModel, Field

class AssistantResponseSchema(BaseModel):
    reply: str = Field(description="The response to the user's question, formatted in markdown.")

class AtsFriendlyTextSchema(BaseModel):
    ats_friendly_description: str = Field(description="The rewritten description using action verbs, clear impact, and ATS-friendly phrasing.")

class AiAssistantView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        message = request.data.get('message')
        resume_version_id = request.data.get('resume_version_id')
        custom_resume_text = request.data.get('custom_resume_text')
        section = request.data.get('section', 'General')

        if not message:
            return Response({"error": "message is required"}, status=status.HTTP_400_BAD_REQUEST)

        resume_text = custom_resume_text or ""
        if resume_version_id:
            try:
                resume_version = ResumeVersion.objects.get(id=resume_version_id, resume__user=request.user)
                resume_text = ""
                for sec in resume_version.sections.all():
                    resume_text += f"{sec.title}:\n{json.dumps(sec.content)}\n\n"
            except ResumeVersion.DoesNotExist:
                return Response({"error": "Resume version not found"}, status=status.HTTP_404_NOT_FOUND)

        try:
            from .config import get_ai_provider
            provider = get_ai_provider()
            system_prompt = "You are a professional career assistant. Answer the user's questions about resume improvement, tailoring, or ATS compatibility."
            prompt = f"User Question: {message}\nResume Section: {section}\nResume:\n{resume_text}"
            
            ai_response = provider.generate_structured_json(prompt, AssistantResponseSchema, system_prompt)
            return Response({"reply": ai_response.reply}, status=status.HTTP_200_OK)
        except Exception as e:
            logger.error(f"AI Assistant Error: {e}")
            return Response({"error": "AI Assistant failed to reply", "details": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class AtsTransformView(APIView):
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser]

    def post(self, request, *args, **kwargs):
        file_obj = request.FILES.get('file')
        custom_resume_text = request.data.get('custom_resume_text')

        if not file_obj and not custom_resume_text:
            return Response({"error": "file or custom_resume_text is required"}, status=status.HTTP_400_BAD_REQUEST)

        raw_text = custom_resume_text or ""
        if file_obj:
            if file_obj.size > 5 * 1024 * 1024:
                return Response({"error": "File size exceeds 5MB limit"}, status=status.HTTP_400_BAD_REQUEST)
            filename = file_obj.name.lower()
            file_bytes = file_obj.read()
            try:
                if filename.endswith('.pdf'):
                    raw_text = extract_text_from_pdf(file_bytes)
                elif filename.endswith('.docx'):
                    raw_text = extract_text_from_docx(file_bytes)
                else:
                    return Response({"error": "Unsupported file format. Please upload PDF or DOCX."}, status=status.HTTP_400_BAD_REQUEST)
            except Exception as e:
                logger.error(f"File parsing error: {e}")
                return Response({"error": "Failed to read the file."}, status=status.HTTP_400_BAD_REQUEST)

        if not raw_text.strip():
            return Response({"error": "No text could be extracted."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            # 1. Extract structured data
            extracted = extract_resume_data(raw_text)

            # 2. Automatically transform experiences to be ATS friendly using AI rewrite
            from .config import get_ai_provider
            provider = get_ai_provider()
            
            for exp in extracted.experience:
                if exp.description.strip():
                    system_prompt = "You are an ATS optimization expert. Rewrite the job experience description to use standard resume action verbs and clear bullet points. Eliminate graphics, jargon, and keep the formatting strictly clean."
                    p = f"Original Description:\n{exp.description}"
                    try:
                        res = provider.generate_structured_json(p, AtsFriendlyTextSchema, system_prompt)
                        from .providers.mock_provider import MockAIProvider
                        if not isinstance(provider, MockAIProvider):
                            exp.description = res.ats_friendly_description
                    except Exception:
                        pass

            # 3. Map to document generator schema
            first_name = ""
            last_name = ""
            if extracted.personal_info.name:
                name_parts = extracted.personal_info.name.split(" ", 1)
                first_name = name_parts[0]
                last_name = name_parts[1] if len(name_parts) > 1 else ""

            mapped_data = {
                "personalInfo": {
                    "firstName": first_name,
                    "lastName": last_name,
                    "email": extracted.personal_info.email or "",
                    "phone": extracted.personal_info.phone or "",
                    "location": extracted.personal_info.location or "",
                    "summary": extracted.summary or ""
                },
                "experience": [
                    {
                        "title": exp.title,
                        "company": exp.company,
                        "startDate": exp.start_date or "",
                        "endDate": exp.end_date or "Present",
                        "description": exp.description
                    } for exp in extracted.experience
                ],
                "education": [
                    {
                        "degree": f"{edu.degree} in {edu.field}" if edu.field else edu.degree,
                        "school": edu.institution,
                        "startDate": edu.start_date or "",
                        "endDate": edu.end_date or "Present"
                    } for edu in extracted.education
                ],
                "skills": extracted.skills,
                "projects": [
                    {
                        "name": proj.name,
                        "description": proj.description
                    } for proj in extracted.projects
                ],
                "certifications": [cert.name for cert in extracted.certifications]
            }

            # 4. Compile to standard DOCX from scratch (ATS friendly layout)
            import docx
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io
            from django.core.files.base import ContentFile
            from documents.models import GeneratedDocument
            from templates.models import Template

            doc = docx.Document()
            
            # Margins
            for sec in doc.sections:
                sec.top_margin = Inches(0.8)
                sec.bottom_margin = Inches(0.8)
                sec.left_margin = Inches(0.8)
                sec.right_margin = Inches(0.8)
                
            # Normal Font Style
            normal_style = doc.styles['Normal']
            normal_font = normal_style.font
            normal_font.name = 'Calibri'
            normal_font.size = Pt(11)
            normal_font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            
            # Header Info (Name, Contact Details)
            p_name = doc.add_paragraph()
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_name = p_name.add_run(f"{first_name} {last_name}".upper())
            run_name.bold = True
            run_name.font.size = Pt(20)
            run_name.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
            
            p_contact = doc.add_paragraph()
            p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_parts = []
            if mapped_data['personalInfo'].get('email'):
                contact_parts.append(mapped_data['personalInfo']['email'])
            if mapped_data['personalInfo'].get('phone'):
                contact_parts.append(mapped_data['personalInfo']['phone'])
            if mapped_data['personalInfo'].get('location'):
                contact_parts.append(mapped_data['personalInfo']['location'])
            p_contact.add_run("  |  ".join(contact_parts))
            p_contact.paragraph_format.space_after = Pt(16)
            
            # Divider and Header function
            def add_section_header(title_str):
                p_hdr = doc.add_paragraph()
                p_hdr.paragraph_format.space_before = Pt(12)
                p_hdr.paragraph_format.space_after = Pt(2)
                p_hdr.paragraph_format.keep_with_next = True
                run_hdr = p_hdr.add_run(title_str.upper())
                run_hdr.bold = True
                run_hdr.font.size = Pt(12)
                run_hdr.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
                
                # Bottom border line
                p_line = doc.add_paragraph()
                p_line.paragraph_format.space_after = Pt(6)
                run_line = p_line.add_run("__________________________________________________________________")
                run_line.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
                run_line.font.size = Pt(6)

            # Summary
            if mapped_data['personalInfo'].get('summary'):
                add_section_header("Professional Summary")
                doc.add_paragraph(mapped_data['personalInfo']['summary'])
                
            # Experience
            if mapped_data.get('experience'):
                add_section_header("Work Experience")
                for exp in mapped_data['experience']:
                    p_job = doc.add_paragraph()
                    p_job.paragraph_format.space_after = Pt(2)
                    p_job.paragraph_format.keep_with_next = True
                    
                    run_title = p_job.add_run(exp['title'])
                    run_title.bold = True
                    
                    p_job.add_run(f"  •  {exp['company']}")
                    
                    dates_str = f"  ({exp['startDate']} - {exp['endDate']})"
                    run_dates = p_job.add_run(dates_str)
                    run_dates.italic = True
                    
                    desc = exp.get('description', '')
                    if desc:
                        for line in desc.split('\n'):
                            line_clean = line.strip()
                            if not line_clean:
                                continue
                            if line_clean.startswith('•') or line_clean.startswith('-') or line_clean.startswith('*'):
                                line_clean = line_clean[1:].strip()
                            p_bullet = doc.add_paragraph(style='List Bullet')
                            p_bullet.add_run(line_clean)
                            p_bullet.paragraph_format.space_after = Pt(2)
                            
            # Education
            if mapped_data.get('education'):
                add_section_header("Education")
                for edu in mapped_data['education']:
                    p_edu = doc.add_paragraph()
                    p_edu.paragraph_format.space_after = Pt(4)
                    run_deg = p_edu.add_run(edu['degree'])
                    run_deg.bold = True
                    p_edu.add_run(f"  •  {edu['school']}")
                    if edu.get('startDate') or edu.get('endDate'):
                        dates_str = f"  ({edu.get('startDate', '')} - {edu.get('endDate', 'Present')})"
                        run_edu_dates = p_edu.add_run(dates_str)
                        run_edu_dates.italic = True

            # Skills
            if mapped_data.get('skills'):
                add_section_header("Skills")
                doc.add_paragraph(", ".join(mapped_data['skills']))
                
            # Projects
            if mapped_data.get('projects'):
                add_section_header("Projects")
                for proj in mapped_data['projects']:
                    p_proj = doc.add_paragraph()
                    p_proj.paragraph_format.space_after = Pt(2)
                    run_pname = p_proj.add_run(proj['name'])
                    run_pname.bold = True
                    
                    desc = proj.get('description', '')
                    if desc:
                        p_desc = doc.add_paragraph(desc)
                        p_desc.paragraph_format.space_after = Pt(4)

            # Save
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Find or use dummy template object to satisfy foreign keys if any
            t_obj = Template.objects.first()
            
            gen_doc = GeneratedDocument.objects.create(
                template=t_obj,
                format='docx',
                user=request.user if request.user.is_authenticated else None
            )
            gen_doc.file.save(f"resume_{last_name or 'ats'}_ats.docx", ContentFile(buffer.read()))
            gen_doc.save()

            return Response({
                "message": "CV successfully transformed into ATS-friendly format",
                "id": gen_doc.id,
                "download_url": f"/api/v1/documents/download/{gen_doc.id}/"
            }, status=status.HTTP_200_OK)

        except Exception as e:
            logger.error(f"ATS Transform Error: {e}")
            return Response({"error": "Failed to transform CV", "details": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

